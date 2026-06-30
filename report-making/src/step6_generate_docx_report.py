from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches, Pt, RGBColor


STEP6_BLUEPRINT_JSON = os.getenv("STEP6_BLUEPRINT_JSON", "outputs/step5_report_blueprint.json")
STEP6_OUTPUT_DOCX = os.getenv("STEP6_OUTPUT_DOCX", "outputs/final_report.docx")


def load_json(path: Path) -> Dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def set_cover_style(paragraph, font_size: int, bold: bool = False, color: RGBColor | None = None) -> None:
    for run in paragraph.runs:
        run.font.size = Pt(font_size)
        run.bold = bold
        if color is not None:
            run.font.color.rgb = color


def shade_row(row, fill_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tcPr.append(shd)


def set_table_grid(table) -> None:
    table.style = "Table Grid"


def add_metrics_strip(document: Document, pairs: List[List[Any]]) -> None:
    labels = [str(p[0]) for p in pairs]
    values = [str(p[1]) for p in pairs]

    table = document.add_table(rows=2, cols=len(labels))
    set_table_grid(table)

    for i, lab in enumerate(labels):
        table.cell(0, i).text = lab
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    for i, val in enumerate(values):
        table.cell(1, i).text = val
        for run in table.cell(1, i).paragraphs[0].runs:
            run.font.size = Pt(10)

    shade_row(table.rows[0], "EFEFEF")

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()


def add_table(document: Document, headers: List[str], rows: List[List[Any]]) -> None:
    cols = len(headers)
    table = document.add_table(rows=1, cols=cols)
    set_table_grid(table)

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    shade_row(table.rows[0], "EFEFEF")

    for r in rows:
        row_cells = table.add_row().cells
        for i in range(cols):
            row_cells[i].text = str(r[i]) if i < len(r) else ""
            for run in row_cells[i].paragraphs[0].runs:
                run.font.size = Pt(9)

    document.add_paragraph()


def add_image(document: Document, img_path: Path, caption: str = "") -> None:
    if not img_path.exists():
        p = document.add_paragraph(f"[Missing image: {img_path}]")
        p.runs[0].font.color.rgb = RGBColor(120, 120, 120)
        return

    section = document.sections[0]
    available_width = section.page_width - section.left_margin - section.right_margin

    document.add_picture(str(img_path), width=available_width)

    if caption:
        p = document.add_paragraph(caption)
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(120, 120, 120)

    document.add_paragraph()


def add_static_toc(document: Document, blueprint: Dict[str, Any]) -> None:
    """
    Static (no-field) TOC: no popup, no update needed.
    Adds dot leaders and (optional) page number column left blank.
    """
    section = document.sections[0]
    avail_width = section.page_width - section.left_margin - section.right_margin

    # Title
    p = document.add_paragraph("Contents")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(16)

    # Right tab stop with dots
    # Place tab stop near right edge of content area
    tab_pos = avail_width

    def toc_line(text: str, level: int = 0, page: str = ""):
        para = document.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25 * level)
        para.paragraph_format.tab_stops.add_tab_stop(
            tab_pos,
            alignment=WD_TAB_ALIGNMENT.RIGHT,
            leader=WD_TAB_LEADER.DOTS
        )
        run = para.add_run(f"{text}\t{page}".rstrip())
        run.font.size = Pt(11)

    sections = blueprint.get("sections", []) or []

    # Show only: section headings (level 0) + heading blocks inside sections (level 1)
    for sec in sections:
        sid = str(sec.get("id", "") or "")
        st = str(sec.get("title", "") or "")
        toc_line(f"{sid}. {st}", level=0, page="")  # page blank to avoid wrong numbers

        # include subheadings from blocks
        blocks = sec.get("blocks", []) or []
        for b in blocks:
            if b.get("type") == "heading":
                lvl = int(b.get("level", 2))
                # only include level-2 headings in TOC (subsections)
                if lvl == 2:
                    toc_line(str(b.get("text", "")), level=1, page="")

    document.add_page_break()


def main() -> None:
    blueprint_path = Path(STEP6_BLUEPRINT_JSON)
    if not blueprint_path.exists():
        raise FileNotFoundError(f"Blueprint JSON not found: {blueprint_path.resolve()}")

    blueprint = load_json(blueprint_path)
    out_docx = Path(STEP6_OUTPUT_DOCX)
    out_docx.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Page setup (A4-ish)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    # -------------------------
    # Cover page
    # -------------------------
    cover = blueprint.get("cover", {}) or {}
    title = str(cover.get("title", "") or "")
    subtitle = str(cover.get("subtitle", "") or "")
    date = str(cover.get("date", "") or "")
    keyword = str(cover.get("keyword", "") or "")
    window = cover.get("window", {}) or {}
    totals = cover.get("totals", {}) or {}

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)
    set_cover_style(p, 26, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(subtitle)
    set_cover_style(p, 12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Date: {date}")
    set_cover_style(p, 11)

    if keyword:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f'Keyword: "{keyword}"')
        set_cover_style(p, 11)

    if window:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Window: {window.get('start_date','')} to {window.get('end_date','')}")
        set_cover_style(p, 11)

    if totals:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Speeches: {totals.get('speeches','')} • Mentions: {totals.get('mentions','')}")
        set_cover_style(p, 11)

    doc.add_page_break()

    # -------------------------
    # Contents (STATIC, compulsory)
    # -------------------------
    add_static_toc(doc, blueprint)

    # -------------------------
    # Sections
    # -------------------------
    assets_root = blueprint_path.parent.parent

    sections = blueprint.get("sections", []) or []
    for sec in sections:
        sec_id = str(sec.get("id", "") or "")
        sec_title = str(sec.get("title", "") or "")
        doc.add_heading(f"{sec_id}. {sec_title}", level=1)

        blocks = sec.get("blocks", []) or []
        for b in blocks:
            btype = b.get("type")

            if btype == "heading":
                lvl = int(b.get("level", 2))
                text = str(b.get("text", "") or "")
                doc.add_heading(text, level=2 if lvl <= 2 else 3)

            elif btype == "paragraph":
                doc.add_paragraph(str(b.get("text", "") or ""))

            elif btype == "bullets":
                for it in (b.get("items", []) or []):
                    doc.add_paragraph(str(it), style="List Bullet")

            elif btype == "metrics_strip":
                add_metrics_strip(doc, b.get("pairs", []) or [])

            elif btype == "table":
                add_table(doc, b.get("headers", []) or [], b.get("rows", []) or [])

            elif btype == "image":
                rel = str(b.get("path", "") or "")
                cap = str(b.get("caption", "") or "")
                pth = Path(rel)
                if not pth.is_absolute():
                    pth = assets_root / pth
                add_image(doc, pth, cap)

        doc.add_page_break()

    doc.save(out_docx)
    print("[DONE] DOCX generated:", out_docx.resolve())


if __name__ == "__main__":
    main()
